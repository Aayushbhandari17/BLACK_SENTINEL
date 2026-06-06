import docx
import os

def extract_text_from_docx(file_path):
    """Extracts text content from paragraphs and tables inside a DOCX file."""
    text_content = []
    try:
        if not os.path.exists(file_path):
            return ""
            
        doc = docx.Document(file_path)
        
        # 1. Extract text from standard paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
                
        # 2. Extract text from tables embedded inside the document
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)
                        
        return "\n".join(text_content)
        
    except Exception as e:
        print(f"[-] DOCX Parsing failed for {file_path}: {str(e)}")
        return ""
    
class DOCXParser:
    def parse(self, file_path):
        return extract_text_from_docx(file_path)